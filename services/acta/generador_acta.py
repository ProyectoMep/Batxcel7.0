"""Genera la nueva acta ABRIENDO Y EDITANDO la última acta generada
(igual que hacía el sistema original con su plantilla): así se conserva
exactamente el diseño del documento (fuente, márgenes, encabezados,
logo insertado a mano, etc.), solo se actualizan los datos que cambian
semana a semana. Solo se construye un documento desde cero la primera
vez que se genera un acta para una entidad (cuando no hay ninguna
anterior que editar).
"""
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

from config.settings import OUTPUT_DIR, BASE_DIR
from models.asistente import AsistenteRepository
from models.observacion import ObservacionRepository
from models.entidad import EntidadRepository
from models.fila_acta import FilaActaRepository
from services.acta.tendencias import calcular_tendencias
from services.acta.extraer_acta import buscar_ultima_acta
from services.acta.estructura_docx import (tabla_info_general, tabla_participantes,
                                           tabla_metricas, tabla_observaciones,
                                           tabla_glosario, limpiar_filas_datos)

DIAS = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]
MESES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
         "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

ROJO_SANTANDER = RGBColor(0xEC, 0x00, 0x00)
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CLARO = "F2F2F2"
FONDO_ENCABEZADO = "1B2A4A"
LOGO = BASE_DIR / "assets" / "logo_santander_color.png"

FUENTE = "Arial"
HORA_REUNION = "3:00 pm – 4:00 pm"   # horario fijo, no se pide en el formulario


def _fecha_es(fecha_iso: str) -> str:
    try:
        d = date.fromisoformat(fecha_iso)
        return f"{DIAS[d.weekday()]} {d.day} de {MESES[d.month]} de {d.year}"
    except Exception:
        return fecha_iso


def _sombrear_celda(celda, color_hex: str):
    tc_pr = celda._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
    tc_pr.append(shd)


def _celda_texto(celda, texto, negrita=False, color=None, tamano=10):
    celda.text = ""
    p = celda.paragraphs[0]
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(tamano)
    run.font.name = FUENTE
    if color is not None:
        run.font.color.rgb = color


def _establecer_fuente_documento(doc: Document):
    """Fuerza Arial como fuente por defecto del documento (estilo
    'Normal'), para que cualquier texto que no tenga fuente explícita
    también salga en Arial."""
    try:
        estilo = doc.styles["Normal"]
        estilo.font.name = FUENTE
        # También hay que setearlo a nivel XML para que aplique en todos
        # los idiomas/plataformas (Word a veces ignora font.name solo).
        rpr = estilo.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FUENTE)
        rfonts.set(qn("w:hAnsi"), FUENTE)
        rfonts.set(qn("w:eastAsia"), FUENTE)
    except Exception:
        pass


def _actualizar_celda_valor(tabla, etiqueta: str, nuevo_valor: str):
    """En una tabla de 2 columnas tipo 'Etiqueta | Valor', busca la fila
    cuya primera celda coincide con 'etiqueta' y reemplaza la segunda."""
    for row in tabla.rows:
        if row.cells[0].text.strip().lower() == etiqueta.lower():
            _celda_texto(row.cells[1], nuevo_valor)
            return True
    return False


def _llenar_participantes(tabla, asistentes):
    limpiar_filas_datos(tabla)
    if not asistentes:
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], "Sin asistentes registrados.")
        return
    for a in asistentes:
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], a.nombre)
        _celda_texto(fila.cells[1], a.cargo)
        # La tabla de participantes SIEMPRE tiene 3 columnas (Nombre,
        # Área/Cargo, Estado) — se escribe directo, sin condicional, para
        # que "Estado" nunca quede vacío por accidente.
        try:
            _celda_texto(fila.cells[2], a.estado or "Asistió")
        except IndexError:
            pass


def _llenar_metricas(tabla, tendencias):
    limpiar_filas_datos(tabla)
    for t in tendencias:
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], t["metrica_acta"])
        _celda_texto(fila.cells[1], t["tendencia"], negrita=True, color=RGBColor(*t["color"]))
        _celda_texto(fila.cells[2], t["descripcion"], tamano=9)
        _celda_texto(fila.cells[3], t["accion"], tamano=9)


def _llenar_observaciones(tabla, observaciones):
    limpiar_filas_datos(tabla)
    if not observaciones:
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], "Sin observaciones registradas.")
        return
    for o in observaciones:
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], f"• {o.texto}")


def _llenar_glosario(tabla, filas_glosario):
    limpiar_filas_datos(tabla)
    for f in filas_glosario:
        if not f.glosario:
            continue
        fila = tabla.add_row()
        _celda_texto(fila.cells[0], f.metrica_acta, negrita=True, tamano=9)
        _celda_texto(fila.cells[1], f.glosario, tamano=9)


def _construir_desde_cero(entidad, tendencias, asistentes, observaciones, filas_glosario,
                          fecha_texto, numero) -> Document:
    """Solo se usa la PRIMERA vez que se genera un acta para una entidad
    (cuando no hay ninguna anterior que editar)."""
    doc = Document()
    _establecer_fuente_documento(doc)

    if LOGO.exists():
        p_logo = doc.add_paragraph()
        run_logo = p_logo.add_run()
        try:
            run_logo.add_picture(str(LOGO), height=Inches(0.4))
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run(f"Acta de Seguimiento (Workstations) — {entidad.nombre}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = FUENTE
    run.font.color.rgb = ROJO_SANTANDER

    def _titulo(numero_sec, texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{numero_sec}.  {texto}")
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = FUENTE
        r.font.color.rgb = NAVY

    _titulo("1", "Información General de la Reunión")
    tabla_info = doc.add_table(rows=4, cols=2)
    tabla_info.style = "Table Grid"
    filas_info = [("Fecha", fecha_texto), ("Hora", HORA_REUNION),
                 ("Tema", entidad.tema_acta), ("Objetivo", entidad.objetivo_acta)]
    for i, (etq, val) in enumerate(filas_info):
        _celda_texto(tabla_info.rows[i].cells[0], etq, negrita=True)
        _sombrear_celda(tabla_info.rows[i].cells[0], GRIS_CLARO)
        _celda_texto(tabla_info.rows[i].cells[1], val)

    p_num = doc.add_paragraph()
    r1 = p_num.add_run("Acta N.º: ")
    r1.bold = True
    r1.font.name = FUENTE
    r2 = p_num.add_run(str(numero))
    r2.font.name = FUENTE
    doc.add_paragraph()

    _titulo("2", "Participantes")
    tp = doc.add_table(rows=1, cols=3)
    tp.style = "Table Grid"
    for i, txt in enumerate(["Nombre", "Área / Cargo", "Estado"]):
        _celda_texto(tp.rows[0].cells[i], txt, negrita=True, color=BLANCO)
        _sombrear_celda(tp.rows[0].cells[i], FONDO_ENCABEZADO)
    _llenar_participantes(tp, asistentes)
    doc.add_paragraph()

    _titulo("3", "Métricas")
    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run("A continuación se presenta el resumen de las métricas y su "
                              "comportamiento, así mismo las respectivas acciones a desarrollar.")
    r_intro.font.name = FUENTE

    tl = doc.add_table(rows=1, cols=3)
    tl.style = "Table Grid"
    leyenda = [("▼ Disminuyó  →  Mejoró", (0x2E, 0x7D, 0x32)),
              ("▲ Aumentó  →  Empeoró", (0xC6, 0x28, 0x28)),
              ("● Se mantuvo  →  Sin cambio", (0x55, 0x55, 0x55))]
    for i, (txt, color) in enumerate(leyenda):
        _celda_texto(tl.rows[0].cells[i], txt, negrita=True, color=RGBColor(*color), tamano=9)
    doc.add_paragraph()

    tm = doc.add_table(rows=1, cols=4)
    tm.style = "Table Grid"
    for i, txt in enumerate(["Métrica", "Tendencia", "Descripción", "Acción Requerida"]):
        _celda_texto(tm.rows[0].cells[i], txt, negrita=True, color=BLANCO)
        _sombrear_celda(tm.rows[0].cells[i], FONDO_ENCABEZADO)
    _llenar_metricas(tm, tendencias)
    doc.add_paragraph()

    _titulo("4", "Observaciones y Decisiones")
    to = doc.add_table(rows=1, cols=1)
    to.style = "Table Grid"
    _llenar_observaciones(to, observaciones)
    doc.add_paragraph()

    _titulo("6", "Glosario de Métricas")
    tg = doc.add_table(rows=1, cols=2)
    tg.style = "Table Grid"
    for i, txt in enumerate(["Métrica", "Descripción"]):
        _celda_texto(tg.rows[0].cells[i], txt, negrita=True, color=BLANCO)
        _sombrear_celda(tg.rows[0].cells[i], FONDO_ENCABEZADO)
    _llenar_glosario(tg, filas_glosario)

    return doc


def generar_acta(entidad_nombre: str, fecha_actual: str, fecha_anterior: str,
                 numero: str, fecha_reunion: str, nombre: str = None) -> dict:
    entidad_nombre = entidad_nombre.upper()
    entidad = EntidadRepository().obtener_por_nombre(entidad_nombre)
    if entidad is None:
        return {"ok": False, "error": f"No existe la entidad '{entidad_nombre}'"}

    tendencias = calcular_tendencias(entidad_nombre, fecha_actual, fecha_anterior)
    if not tendencias:
        return {"ok": False,
                "error": "No se pudo calcular ninguna tendencia. Revisa que existan los reportes "
                         "de ambas fechas y que haya filas de acta configuradas."}

    asistentes = AsistenteRepository(entidad_nombre).listar()
    observaciones = ObservacionRepository(entidad_nombre).listar()
    filas_glosario = FilaActaRepository().ordenadas()
    fecha_texto = _fecha_es(fecha_reunion)

    ruta_base = buscar_ultima_acta(entidad_nombre)

    if ruta_base is not None:
        # ── Camino normal: abrir y editar la última acta generada ──
        doc = Document(ruta_base)
        _establecer_fuente_documento(doc)

        t_info = tabla_info_general(doc)
        if t_info is not None:
            _actualizar_celda_valor(t_info, "Fecha", fecha_texto)
            _actualizar_celda_valor(t_info, "Hora", HORA_REUNION)
            _actualizar_celda_valor(t_info, "Tema", entidad.tema_acta)
            _actualizar_celda_valor(t_info, "Objetivo", entidad.objetivo_acta)

        # Actualizar "Acta N.º:" buscándolo en los párrafos
        for p in doc.paragraphs:
            if "acta n" in p.text.lower() and ":" in p.text:
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = f"Acta N.º: {numero}"
                    p.runs[0].font.name = FUENTE
                else:
                    r = p.add_run(f"Acta N.º: {numero}")
                    r.font.name = FUENTE
                break

        t_part = tabla_participantes(doc)
        if t_part is not None:
            _llenar_participantes(t_part, asistentes)

        t_met = tabla_metricas(doc)
        if t_met is not None:
            _llenar_metricas(t_met, tendencias)

        t_obs = tabla_observaciones(doc)
        if t_obs is not None:
            _llenar_observaciones(t_obs, observaciones)

        # El glosario NO se toca — se conserva tal cual estaba, ya que no
        # cambia semana a semana (a menos que agregues métricas nuevas).
    else:
        # ── Primera vez para esta entidad: construir desde cero ──
        doc = _construir_desde_cero(entidad, tendencias, asistentes, observaciones,
                                    filas_glosario, fecha_texto, numero)

    nombre = (nombre or f"Acta {numero} de seguimiento (Workstations) {entidad_nombre}").strip()
    nombre = re.sub(r'[\\/:*?"<>|]+', "_", nombre)
    carpeta = OUTPUT_DIR / fecha_actual / "actas" / entidad_nombre
    carpeta.mkdir(parents=True, exist_ok=True)
    ruta_salida = carpeta / f"{nombre}.docx"
    doc.save(ruta_salida)

    return {"ok": True, "archivo": ruta_salida.name, "fecha": fecha_actual,
            "entidad": entidad_nombre, "metricas": len(tendencias),
            "asistentes": len(asistentes), "observaciones": len(observaciones),
            "editado_desde": ruta_base.name if ruta_base else None}


def listar_actas() -> dict:
    entidades = EntidadRepository().listar()
    resultado = {e.nombre: [] for e in entidades}
    if not OUTPUT_DIR.exists():
        return resultado
    for carpeta in sorted(OUTPUT_DIR.iterdir(), reverse=True):
        if not carpeta.is_dir():
            continue
        for entidad in entidades:
            sub = carpeta / "actas" / entidad.nombre
            if sub.is_dir():
                for f in sorted(sub.glob("*.docx")):
                    resultado[entidad.nombre].append({
                        "fecha": carpeta.name, "nombre": f.stem,
                        "docx": f"actas/{entidad.nombre}/{f.name}"})
    return resultado